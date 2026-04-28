from __future__ import annotations
import re
import httpx
import pdfplumber
import io
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from services.extraction_constants import (
    CONTROL_KEYWORDS,
    DOMAIN_MAP,
    DETAIL_KEYWORDS,
    IGNORE_PATTERNS
)


async def extract_controls_from_url(url: str) -> list:
    """Scrape a trust center URL and extract controls using local parsing and heuristic-based analysis."""
    text = await _scrape_with_playwright(url)

    if not text or len(text.strip()) < 100:
        # Fallback to httpx for server-rendered pages
        text = await _scrape_with_httpx(url)

    if not text or len(text.strip()) < 100:
        raise ValueError(f"Could not extract meaningful content from {url}")

    return _extract_controls_locally(text, prefix="TC")


async def extract_controls_from_pdf(content: bytes, filename: str) -> list:
    """Parse a PDF and extract controls using local parsing and heuristic-based analysis."""
    text = _parse_pdf(content)
    if not text or len(text.strip()) < 100:
        raise ValueError(f"Could not extract text from PDF: {filename}")
    return _extract_controls_locally(text, prefix="DOC")


async def extract_controls_from_docx(content: bytes, filename: str) -> list:
    """Parse a DOCX and extract controls using local parsing and heuristic-based analysis."""
    text = _parse_docx(content)
    if not text or len(text.strip()) < 100:
        raise ValueError(f"Could not extract text from DOCX: {filename}")
    return _extract_controls_locally(text, prefix="DOC")


async def _scrape_with_playwright(url: str) -> str:
    """Use Playwright to render JS-heavy SPAs — returns inner_text directly."""
    try:
        from playwright.async_api import async_playwright

        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) "
                           "Chrome/120.0.0.0 Safari/537.36"
            )
            page = await context.new_page()
            await page.goto(url, wait_until="networkidle", timeout=30000)
            await page.wait_for_timeout(3000)

            # Scroll to trigger lazy-loaded content
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(1500)
            await page.evaluate("window.scrollTo(0, 0)")
            await page.wait_for_timeout(500)

            # Use inner_text — gives clean rendered text, no HTML parsing needed
            text = await page.inner_text("body")
            await browser.close()

        # Clean up whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        print(f"[Playwright] Extracted {len(text)} chars from {url}")
        return text[:20000]

    except Exception as e:
        print(f"[Playwright] Failed: {e}, falling back to httpx")
        return ""


async def _scrape_with_httpx(url: str) -> str:
    """Fallback for server-rendered pages."""
    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; ControlMapper/1.0)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    async with httpx.AsyncClient(timeout=30, follow_redirects=True) as http:
        resp = await http.get(url, headers=headers)
        resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text[:20000]


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)[:20000]


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX including tables."""
    doc = Document(BytesIO(content))
    text_parts = []
    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)
    return "\n\n".join(text_parts)[:20000]


def _extract_controls_locally(text: str, prefix: str) -> list[dict]:
    """
    Heuristic-based control extraction.
    """
    controls = []
    lines = text.split("\n")
    idx = 1
    seen = set()
    for raw_line in lines:
        line = raw_line.strip()
        if not _is_valid_control_candidate(line):
            continue
        normalized_line = line.lower()
        if normalized_line in seen:
            continue
        seen.add(normalized_line)
        controls.append({
            "control_id": f"{prefix}-{idx}",
            "text": line,
            "domain": _classify_domain(line),
            "specificity": _compute_specificity(line),
        })
        idx += 1
    return controls


def _is_valid_control_candidate(line: str) -> bool:
    """
    Determine whether a line is likely a security/compliance control.
    """
    if len(line) < 40:
        return False
    lower = line.lower()
    if any(pattern in lower for pattern in IGNORE_PATTERNS):
        return False
    if not any(keyword in lower for keyword in CONTROL_KEYWORDS):
        return False
    return True


def _classify_domain(line: str) -> str:
    """
    Assign a security domain using keyword heuristics.
    """
    lower = line.lower()
    for keyword, domain in DOMAIN_MAP.items():
        if keyword in lower:
            return domain
    return "General"


def _compute_specificity(line: str) -> str:
    """
    Estimate control specificity based on detail indicators.
    """
    lower = line.lower()
    detail_score = sum(
        1
        for keyword in DETAIL_KEYWORDS
        if keyword in lower
    )
    if detail_score >= 2:
        return "high"
    if detail_score == 1:
        return "medium"
    return "low"
